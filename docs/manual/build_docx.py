"""
Battery Pack Intelligence 사용자 매뉴얼 → Word(.docx) 변환 스크립트
실행: python docs/manual/build_docx.py
출력: docs/manual/user-manual.docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = os.path.dirname(os.path.abspath(__file__))
IMG = os.path.join(BASE, "images")

# ──────────────────────────────────────────
# 헬퍼
# ──────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run("※ " + text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x44, 0x88, 0xCC)
    run.italic = True
    return p

def add_image(doc, filename, width_inches=5.5, caption=None):
    path = os.path.join(IMG, filename)
    if not os.path.exists(path):
        doc.add_paragraph(f"[이미지 없음: {filename}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.size = Pt(9)
        cp.runs[0].italic = True
        cp.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 헤더
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(hdr[i], "1F3864")
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # 데이터
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
            row[i].paragraphs[0].runs[0].font.size = Pt(9)
    # 열 너비
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    # 배경색 (연한 회색)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p

def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ──────────────────────────────────────────
# 문서 생성
# ──────────────────────────────────────────

doc = Document()

# 기본 여백 설정
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# 기본 폰트
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(10)

# ── 표지 ──────────────────────────────────
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Battery Pack Intelligence")
tr.bold = True
tr.font.size = Pt(22)
tr.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("사용자 매뉴얼")
sr.font.size = Pt(16)
sr.bold = True

doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
ver.add_run("SL Power 니켈 플레이트 설계 검증기  v0.2.25").font.size = Pt(11)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run("2026-04-23").font.size = Pt(10)

doc.add_paragraph()
add_image(doc, "01_overview.png", width_inches=5.8, caption="그림 1 — 프로그램 초기 화면")
doc.add_page_break()

# ── 1. 프로그램 개요 ──────────────────────
add_heading(doc, "1. 프로그램 개요", 1)
add_para(doc, "Battery Pack Intelligence는 SL Power의 니켈 플레이트 설계 검증 도구입니다. "
              "배터리 셀의 직렬·병렬 조합(S×P), 홀더 형상, 니켈 플레이트 패턴을 시각적으로 탐색하고 검증합니다.")
doc.add_paragraph()

add_table(doc,
    ["기능", "설명"],
    [
        ["셀 배열 렌더링",   "정배열·엇배열·커스텀 배열을 SVG로 정밀 시각화"],
        ["그룹 배정 탐색",   "홀더 제약을 만족하는 셀→직렬그룹 배정 자동 열거"],
        ["부분 고정 탐색",   "일부 그룹을 수동 고정 후 나머지만 자동 탐색"],
        ["니켈 플레이트 패턴", "상면·하면 니켈 연결 패턴 자동 생성 및 검증"],
        ["비용 모델",         "m_distinct(금형 종류 수) 기반 금형비 추정"],
        ["SVG / PNG 내보내기", "설계 결과를 파일로 출력"],
    ],
    col_widths=[4.5, 10.5]
)

doc.add_paragraph()
add_note(doc, "실행 방법: battery_pack_renderer.html을 Chrome 또는 Edge 브라우저에서 여세요.")
add_hr(doc)

# ── 2. 화면 구성 ──────────────────────────
add_heading(doc, "2. 화면 구성", 1)
add_para(doc, "화면은 좌측 패널(설정), 중앙 캔버스(SVG 렌더링), 우측 패널(후보 목록) 3개 영역으로 구성됩니다.")
doc.add_paragraph()
add_image(doc, "02_generated_result.png", width_inches=5.8, caption="그림 2 — 레이아웃 생성 후 전체 화면 (3S3P, 7개 후보)")
doc.add_paragraph()

add_table(doc,
    ["상단 바 요소", "설명"],
    [
        ["Both Faces / Top Face / Bottom Face", "표시할 면 선택"],
        ["− / 100% / +",                        "캔버스 확대·축소"],
        ["우측 상태 텍스트",                     "현재 설정 요약 (셀 규격·S×P·셀수·BMS방향·m값)"],
    ],
    col_widths=[6, 9]
)
add_hr(doc)

# ── 3. 기본 설정 ──────────────────────────
add_heading(doc, "3. 기본 설정 (좌측 패널)", 1)

add_heading(doc, "3.1 Cell Type", 2)
add_para(doc, "18650 또는 21700 중 하나를 선택합니다. 파란색 배경이 현재 선택된 타입입니다.")
add_table(doc,
    ["셀 타입", "직경", "주요 용도"],
    [
        ["18650", "18 mm", "소형·경량 팩"],
        ["21700", "21 mm", "고용량·고출력 팩"],
    ],
    col_widths=[3, 3, 9]
)

doc.add_paragraph()
add_heading(doc, "3.2 홀더 크기 (원칙 17 L1)", 2)
add_para(doc, "물리적인 배터리 홀더의 행(row) × 열(column) 크기를 설정합니다. "
              "auto로 두면 S×P 입력값에서 자동 계산됩니다. − / + 버튼으로 행·열 수를 조정합니다.")
add_note(doc, "예: 13×4 홀더에 52셀 배치 → 공란 없음. 13×5 홀더에 52셀 → 8개 공란 발생.")

doc.add_paragraph()
add_heading(doc, "3.3 직렬(S) × 병렬(P)", 2)
add_para(doc, "배터리 팩의 직렬 그룹 수(S)와 병렬 셀 수(P)를 설정합니다. 총 셀 수 = S × P.")
add_table(doc,
    ["파라미터", "의미", "관계식"],
    [
        ["S (Series)",  "직렬 그룹 수", "팩 전압 = 셀 전압 × S"],
        ["P (Parallel)", "병렬 셀 수",  "팩 용량 = 셀 용량 × P"],
    ],
    col_widths=[3.5, 4.5, 7]
)

doc.add_paragraph()
add_heading(doc, "3.4 Arrangement (배열 방식)", 2)
add_table(doc,
    ["버튼", "배열", "설명"],
    [
        ["정배열", "Square",    "셀이 격자(grid) 형태로 정렬. 홀·짝 행 오프셋 없음"],
        ["엇배열", "Staggered", "홀수 행이 pitch/2 오프셋. 헥사고날 밀집 배치"],
        ["커스텀", "Custom",    "행별로 셀 수를 다르게 지정. 불규칙 형상 팩에 사용"],
    ],
    col_widths=[2.5, 3, 9.5]
)
add_hr(doc)

# ── 4. BMS 위치 설정 ──────────────────────
add_heading(doc, "4. BMS 위치 설정", 1)
add_para(doc, "BMS(배터리 관리 시스템) 커넥터의 위치를 설정합니다. "
              "캔버스에 초록색 삼각형 마커와 BMS 라벨로 표시됩니다.")
add_table(doc,
    ["항목", "설명"],
    [
        ["BMS Position 엣지",  "Top / Bot(기본) / Left / Right 중 선택"],
        ["BMS Position 비율",  "0%=엣지 시작점, 50%=중앙(기본), 100%=끝점"],
    ],
    col_widths=[4.5, 10.5]
)
add_hr(doc)

# ── 5. B+/B- 방향 ─────────────────────────
add_heading(doc, "5. B+ / B− 방향 설정", 1)
add_para(doc, "팩의 양극(B+)과 음극(B−) 출력 단자 방향을 설정합니다.")
add_table(doc,
    ["버튼", "의미"],
    [
        ["Top",   "상단 방향"],
        ["Left",  "좌측 방향 (B+ 기본값)"],
        ["Right", "우측 방향 (B− 기본값)"],
        ["Bot",   "하단 방향"],
    ],
    col_widths=[3, 12]
)
add_note(doc, "원칙 4: B+ 방향과 B− 방향은 서로 반대 엣지여야 합니다. "
              "같은 방향으로 설정하면 탐색 결과가 0개가 될 수 있습니다.")
add_hr(doc)

# ── 6. 고급 설정 ──────────────────────────
add_heading(doc, "6. 고급 설정", 1)

add_heading(doc, "6.1 최대 플레이트 수 (0=제한없음)", 2)
add_para(doc, "금형(m_distinct)의 최대 종류 수를 제한합니다. "
              "0이면 제한 없음. N 설정 시 고유 형상이 N종 초과인 후보는 필터링됩니다.")

doc.add_paragraph()
add_heading(doc, "6.2 1번 셀 위치 (선택)", 2)
add_para(doc, "B+ 그룹(G0)의 첫 번째 셀 위치 힌트를 제공합니다.")
add_table(doc,
    ["옵션", "의미"],
    [
        ["자동 (기본)", "B+ 경계 셀 중 자동 선택"],
        ["TL",          "Top-Left — 좌상단 코너 셀"],
        ["TR",          "Top-Right — 우상단 코너 셀"],
        ["BL",          "Bottom-Left — 좌하단 코너 셀"],
        ["BR",          "Bottom-Right — 우하단 코너 셀"],
    ],
    col_widths=[3.5, 11.5]
)

doc.add_paragraph()
add_heading(doc, "6.3 ICC 제약 조건 (원칙 16)", 2)
add_para(doc, "Industrial Compactness Constraint — 제조 실용성 필터입니다. 토글 ON/OFF로 개별 활성화합니다.")
add_table(doc,
    ["제약", "조건", "의미"],
    [
        ["ICC①", "행 스팬 ≤ 2",    "하나의 그룹이 3행 이상 걸치면 금지"],
        ["ICC②", "종횡비 ≤ 2.0",   "그룹 너비/높이 비율이 2.0 초과면 금지"],
        ["ICC③", "볼록성 ≥ 0.75",  "그룹 면적 / 볼록 껍질 면적 비율"],
    ],
    col_widths=[2.5, 4, 8.5]
)
add_note(doc, "제약을 끄면 더 많은 후보가 탐색되지만, 실제 제조 불가능한 형상이 포함될 수 있습니다.")

doc.add_paragraph()
add_heading(doc, "6.4 Display Options", 2)
add_table(doc,
    ["옵션", "설명"],
    [
        ["Nickel plate",    "니켈 플레이트 연결선 표시/숨기기"],
        ["B+ / B−",         "양·음극 단자 마커 표시/숨기기"],
        ["셀 좌표 (rNcM)",  "각 셀 위에 행·열 좌표 표시. 부분 고정 탐색 입력 시 유용"],
    ],
    col_widths=[4, 11]
)

doc.add_paragraph()
add_heading(doc, "6.5 탐색 시간 제한", 2)
add_table(doc,
    ["버튼", "시간"],
    [
        ["1분",   "60초"],
        ["10분",  "600초 (기본값)"],
        ["30분",  "1800초"],
        ["1시간", "3600초"],
    ],
    col_widths=[3, 12]
)
add_hr(doc)

# ── 7. 커스텀 배열 ────────────────────────
add_heading(doc, "7. 커스텀 배열 모드", 1)
add_image(doc, "09_custom_arrangement.png", width_inches=5.5, caption="그림 3 — 커스텀 배열 모드 활성화 상태")
doc.add_paragraph()
add_para(doc, "Arrangement에서 커스텀 버튼을 클릭하면 행별 셀 수 입력 영역이 나타납니다.")

add_heading(doc, "7.1 행별 셀 수 입력", 2)
add_para(doc, "텍스트 박스에 행별 셀 수를 입력합니다. 한 줄 = 한 행 (위에서 아래 방향, 행 0부터).")
add_code_block(doc, "10\n12\n13\n12\n10")
add_para(doc, "쉼표·콜론 형식도 지원합니다:")
add_code_block(doc, "10,12:-2,13:-2,12:-5,-1")
add_para(doc, "입력 후 S×P와 합계 일치 여부를 실시간 검증합니다.")

add_heading(doc, "7.2 엇배열 (Staggered) 토글", 2)
add_table(doc,
    ["상태", "의미"],
    [
        ["ON (파란색)", "홀수 행 pitch/2 오프셋 적용"],
        ["OFF",         "모든 행 동일 x 시작점"],
    ],
    col_widths=[4, 11]
)
add_hr(doc)

# ── 8. Generate Layout ────────────────────
add_heading(doc, "8. Generate Layout 실행", 1)
add_para(doc, "모든 설정을 완료한 후 파란색 Generate Layout 버튼을 클릭합니다.")
add_table(doc,
    ["단계", "동작"],
    [
        ["1", "홀더 그리드 구성 (buildHolderGrid)"],
        ["2", "B+/B− 경계 셀 계산 (calcBoundarySet)"],
        ["3", "그룹 배정 탐색 시작 (Web Worker 병렬)"],
        ["4", "첫 번째 후보 자동 선택 → 캔버스 렌더링"],
        ["5", "우측 패널에 전체 후보 목록 표시"],
    ],
    col_widths=[1.5, 13.5]
)
add_note(doc, "탐색 중에는 상단 바에 '탐색 중... N개' 형태로 실시간 후보 수가 표시됩니다.")
doc.add_paragraph()
add_note(doc, "후보 0개 표시 시: B+/B− 방향 확인 → 홀더 크기 확인 → ICC 제약 일부 해제 → 시간 제한 증가 순으로 확인하세요.")
add_hr(doc)

# ── 9. 캔버스 뷰 ──────────────────────────
add_heading(doc, "9. 캔버스 뷰", 1)

add_heading(doc, "9.1 면 선택", 2)
add_image(doc, "03_top_face_view.png", width_inches=5.5, caption="그림 4 — Top Face 단독 뷰")
add_image(doc, "04_bottom_face_view.png", width_inches=5.5, caption="그림 5 — Bottom Face 단독 뷰")
add_table(doc,
    ["버튼", "표시 내용"],
    [
        ["Both Faces",   "상면(top face)과 하면(bottom face)을 좌우로 나란히 표시"],
        ["Top Face",     "상면만 단독 표시. B+ 단자가 있는 면"],
        ["Bottom Face",  "하면만 단독 표시. B− 단자가 있는 면"],
    ],
    col_widths=[3.5, 11.5]
)

add_heading(doc, "9.2 SVG 범례", 2)
add_image(doc, "05_zoomed_canvas.png", width_inches=5.5, caption="그림 6 — 확대된 캔버스 (130%)")
add_table(doc,
    ["시각 요소", "의미"],
    [
        ["빨간 테두리 원",   "양극(+) 셀 상면"],
        ["검정 테두리 원",   "음극(−) 셀 상면"],
        ["빨간 중앙 점",     "하면 음극 마커"],
        ["회색 선",          "니켈 플레이트 연결선"],
        ["빨간 박스 + 점",   "B+ 단자 탭"],
        ["파란 박스 + 점",   "B− 단자 탭"],
        ["초록 삼각형",      "BMS 위치 마커"],
        ["rNcM 텍스트",      "셀 좌표 (Display Options ON 시)"],
    ],
    col_widths=[4, 11]
)
add_hr(doc)

# ── 10. 후보 패널 ─────────────────────────
add_heading(doc, "10. 셀 배열 후보 패널 (우측)", 1)
add_image(doc, "11_candidates_panel.png", width_inches=5.5, caption="그림 7 — 후보 카드 목록 (7개)")
doc.add_paragraph()
add_para(doc, "Generate Layout 실행 후 우측 패널에 탐색된 후보 목록이 표시됩니다.")

add_heading(doc, "10.1 후보 카드 읽는 법", 2)
add_code_block(doc, "보스트로페돈 L→R   snake          neutral\n○ ○ ○\nΣ +0 · 플레이트 2종 · 행스팬 평균 1.0/최대 1")
add_table(doc,
    ["요소", "의미"],
    [
        ["배열 이름",       "그룹 배정 방향 전략 (보스트로페돈 L→R / R→L, 열 우선, triomino 등)"],
        ["전략 태그",       "snake / triomino 등 탐색 알고리즘 종류"],
        ["극성 태그",       "neutral / positive / negative"],
        ["○ ○ ○",          "직렬 그룹의 B+ 셀 연결 상태"],
        ["Σ +0",            "품질 점수 (클수록 우수)"],
        ["플레이트 N종",    "m_distinct — 고유 금형 종류 수"],
        ["행스팬 평균/최대", "그룹이 걸치는 행 수"],
    ],
    col_widths=[4, 11]
)
add_note(doc, "카드 클릭 → 해당 후보가 캔버스에 렌더링됩니다. 선택된 카드는 빨간 테두리로 강조됩니다.")

add_heading(doc, "10.2 후보 선택 결과", 2)
add_image(doc, "12_candidate_selected.png", width_inches=5.5, caption="그림 8 — 두 번째 후보 선택 후 캔버스 변경")

add_heading(doc, "10.3 필터", 2)
add_table(doc,
    ["필터", "설명"],
    [
        ["형상종류", "전체 / 2종 / 3종 등 m_distinct 값으로 필터"],
        ["품질 Σ",  "전체 / +0 이상 / +1 이상 등 품질 점수 기준 필터"],
    ],
    col_widths=[3.5, 11.5]
)
add_hr(doc)

# ── 11. 부분 고정 탐색 ───────────────────
add_heading(doc, "11. 부분 고정 탐색", 1)
add_image(doc, "10_pinned_groups.png", width_inches=5.5, caption="그림 9 — 부분 고정 탐색 섹션 (Sparse 기본값 표시)")
doc.add_paragraph()
add_para(doc, "커스텀 배열 모드에서만 사용 가능한 기능입니다. "
              "일부 직렬 그룹의 셀 위치를 수동으로 고정하고, 나머지 그룹만 자동으로 탐색합니다.")

add_heading(doc, "11.1 연속 모드", 2)
add_para(doc, "G0부터 순서대로 각 그룹의 셀 좌표를 한 줄씩 입력합니다.")
add_code_block(doc, "r3c0 r2c0 r1c0 r2c1\nr4c0 r4c1 r3c1 r3c2")
add_note(doc, "좌표 형식: rNcM = N번째 행(row), M번째 열(column). 0-indexed.")

add_heading(doc, "11.2 Sparse(비연속) 모드", 2)
add_para(doc, "임의의 그룹 번호를 직접 지정하여 고정합니다. 비연속 인덱스 동시 고정 가능.")
add_code_block(doc, "0: r3c0 r2c0 r1c0 r2c1\n11: r3c10 r2c11 r1c10 r0c8\n12: r3c11 r2c12 r1c11 r0c9")
add_para(doc, "위 예시는 13S4P 배열에서 G0, G11, G12를 고정하고 G1~G10만 탐색합니다.")
add_note(doc, "셀 좌표 확인: Display Options에서 셀 좌표(rNcM) 토글 ON → 캔버스에 좌표 표시.")

add_heading(doc, "11.3 버튼", 2)
add_table(doc,
    ["버튼", "동작"],
    [
        ["초기화",      "텍스트 박스 내용을 지움"],
        ["📌 고정 탐색", "입력된 고정 그룹으로 부분 탐색 시작"],
    ],
    col_widths=[4, 11]
)
add_hr(doc)

# ── 12. Pack Summary ──────────────────────
add_heading(doc, "12. Pack Summary", 1)
add_image(doc, "08_sidebar_bottom.png", width_inches=5.5, caption="그림 10 — Pack Summary + Export 버튼")
doc.add_paragraph()

add_heading(doc, "12.1 기본 정보", 2)
add_table(doc,
    ["항목", "설명"],
    [
        ["cells",   "총 셀 수 (S×P 표기)"],
        ["nickel",  "니켈 플레이트 수 (상면 + 하면)"],
        ["layout",  "배열 방향 (좌우 / 상하)"],
        ["B+",      "B+ 단자 위치 (행 인덱스)"],
        ["B−",      "B− 단자 위치 (행 인덱스)"],
        ["크기",    "팩 외형 치수 W × H (mm)"],
    ],
    col_widths=[3, 12]
)

add_heading(doc, "12.2 비용 모델", 2)
add_table(doc,
    ["항목", "설명"],
    [
        ["m_min (est.)", "최소 금형 종류 수"],
        ["mold cost",    "금형비 추정 (m_distinct × ₩20M)"],
        ["reuse",        "재사용률 (중복 사용 형상 / 전체 플레이트 수)"],
        ["feasible",     "제조 가능성 검증 결과 (✓ 가능 / ✗ 불가)"],
    ],
    col_widths=[3.5, 11.5]
)

add_heading(doc, "12.3 BMS Wire (원칙 16)", 2)
add_table(doc,
    ["항목", "설명", "색상 기준"],
    [
        ["dist B+",    "BMS → B+ 맨해튼 거리 (mm)",  "녹색 < 60mm / 황색 > 120mm"],
        ["dist B−",    "BMS → B− 맨해튼 거리 (mm)",  "동일"],
        ["total wire", "총 BMS 와이어 길이 (mm)",     "녹색 < 120mm / 황색 > 240mm"],
    ],
    col_widths=[3, 7, 5]
)
add_hr(doc)

# ── 13. 내보내기 ──────────────────────────
add_heading(doc, "13. 파일 내보내기", 1)
add_table(doc,
    ["버튼", "출력 형식", "파일명 예시"],
    [
        ["↓ Export SVG", "벡터 SVG",         "3s3p_21700_square.svg"],
        ["↓ Export PNG", "래스터 PNG (2×)",  "3s3p_21700_square.png"],
    ],
    col_widths=[3.5, 4, 7.5]
)
add_note(doc, "현재 캔버스에 표시된 면 기준으로 내보냅니다. "
              "Both Faces 시 상·하면 모두, Top Face 시 상면만 포함됩니다.")
add_hr(doc)

# ── 14. FAQ ───────────────────────────────
add_heading(doc, "14. 자주 묻는 질문 (FAQ)", 1)

faqs = [
    ("후보가 하나도 탐색되지 않습니다.",
     "① B+/B− 방향이 서로 다른 엣지인지 확인\n"
     "② 홀더 크기가 S×P 셀을 수용할 수 있는지 확인\n"
     "③ ICC 제약 조건 중 하나 이상을 OFF로 변경\n"
     "④ 탐색 시간 제한을 늘려서 재시도"),
    ("탐색이 오래 걸립니다.",
     "대형 팩(S·P 클수록)은 탐색 공간이 급격히 커집니다. "
     "부분 고정 탐색(11장)을 활용하면 탐색 공간을 크게 줄일 수 있습니다."),
    ("커스텀 배열에서 셀 수 불일치 오류가 납니다.",
     "입력한 각 행의 셀 수 합계가 S×P와 일치하지 않습니다. "
     "셀 수를 다시 계산하거나 S / P 값을 조정하세요."),
    ("부분 고정 탐색에서 셀 좌표를 어떻게 확인하나요?",
     "Display Options에서 셀 좌표(rNcM) 토글을 ON 하면 "
     "캔버스의 각 셀 위에 r2c3 형태로 좌표가 표시됩니다."),
    ("Export PNG 파일이 어디에 저장되나요?",
     "브라우저 기본 다운로드 폴더에 저장됩니다. (보통 ~/Downloads)"),
    ("Sparse 형식과 연속 형식은 어떻게 구분되나요?",
     "입력에 '숫자:' 패턴(그룹 번호 + 콜론)이 포함되면 Sparse 모드로 자동 인식합니다. "
     "없으면 연속 모드(G0부터 순서대로)로 처리됩니다."),
]

for q, a in faqs:
    qp = doc.add_paragraph()
    qr = qp.add_run("Q. " + q)
    qr.bold = True
    qr.font.size = Pt(10)
    qr.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    ap = doc.add_paragraph()
    ap.paragraph_format.left_indent = Cm(0.5)
    ar = ap.add_run("A. " + a)
    ar.font.size = Pt(10)
    doc.add_paragraph()

add_hr(doc)

# ── 저작권 표기 ───────────────────────────
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = foot.add_run("Battery Pack Intelligence  ·  SL Power Co., Ltd. · SELA  ·  이론 v0.2.24 / 앱 v0.2.25")
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── 저장 ──────────────────────────────────
out = os.path.join(BASE, "user-manual.docx")
doc.save(out)
print("저장 완료: " + out)
